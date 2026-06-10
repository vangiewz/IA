import os
import json
import httpx
import pandas as pd
from typing import List, Dict, Any
from io import BytesIO

# Importaciones para formatos de archivo
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import zipfile

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.drawing.image import Image as OpenpyxlImage

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from anthropic import AsyncAnthropic

class ReportGeneratorService:
    def __init__(self):
        self.api_key = os.getenv("CLAUDE_API_KEY")
        self.client = AsyncAnthropic(api_key=self.api_key)
        
        backend_host = os.getenv("BACKEND_PROD_URL", "http://backend:8080")
        self.backend_url = f"{backend_host}/api/tramites/report-data"

    async def procesar_chat(self, historial: List[Dict[str, str]]) -> Dict[str, Any]:
        """
        Envía el historial a Claude con un System Prompt estricto.
        Si está completo, llama al backend, genera el archivo y devuelve el stream.
        Si está incompleto, devuelve el JSON de Claude con la pregunta a hacer al usuario.
        """
        system_prompt = """
Eres un asistente de reportes BI. Tu objetivo es extraer 5 pilares del historial del chat:
1. DATOS (qué datos quiere ver el usuario de los trámites).
2. FILTROS. SOLO TIENES PERMITIDO extraer y usar estas llaves EXACTAS: 
   - 'fechaInicio' (YYYY-MM-DD)
   - 'fechaFin' (YYYY-MM-DD)
   - 'nombreCliente'
   - 'departamentoAsignado'
   - 'nombrePlantilla'
   - 'estadoGlobal' (ej. FINALIZADO, PENDIENTE)
3. FORMATOS (Puede ser una lista con uno o varios de: 'EXCEL', 'WORD', 'PDF').
4. AGRUPACION. Si el usuario pide explícitamente estadísticas por "departamentos", "funcionarios" o métricas agrupadas, usa "DEPARTAMENTO" o "FUNCIONARIO". De lo contrario, usa "TRAMITE".
5. GRAFICOS. Si el usuario pide generar gráficos visuales, indica una lista (ej. ["BARRAS", "PASTEL", "LINEAS"]). Si no pide explícitamente, deja la lista vacía [].

Si falta ALGUNO de los pilares 1, 2, 3 o 4, devuelve ESTRICTAMENTE este JSON:
{"estado": "INCOMPLETO", "mensaje_usuario": "Aquí le preguntas o aclaras qué le falta"}

Si tienes los pilares claros y completos, devuelve ESTRICTAMENTE este JSON:
{"estado": "COMPLETO", "parametros_query": {"datos": "descripción", "filtros": {"llave1": "valor1"}, "formatos": ["EXCEL"], "agrupacion": "TRAMITE|DEPARTAMENTO|FUNCIONARIO", "graficos": ["BARRAS"]}}

NO DEBES añadir ningún texto adicional ni markdown. Solo el objeto JSON crudo."""

        messages = []
        for msg in historial:
            role = "assistant" if msg["role"] == "ia" else "user"
            messages.append({"role": role, "content": msg["content"]})

        try:
            response = await self.client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1024,
                system=system_prompt,
                messages=messages
            )

            response_text = response.content[0].text.strip()
            
            # Extraer solo la parte JSON en caso de que la IA añada texto adicional
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}')
            
            if start_idx != -1 and end_idx != -1 and end_idx >= start_idx:
                json_str = response_text[start_idx:end_idx+1]
                response_json = json.loads(json_str)
            else:
                # Si no hay JSON válido, asumimos que es una respuesta de texto directo de Claude
                return {"estado": "INCOMPLETO", "mensaje_usuario": response_text}
            
            if response_json.get("estado") == "COMPLETO":
                # Llama a Spring Boot, genera el archivo
                file_stream, file_name, mime_type = await self._generar_reporte_fisico(response_json)
                return {
                    "estado": "COMPLETO",
                    "file_stream": file_stream,
                    "file_name": file_name,
                    "mime_type": mime_type
                }
            else:
                return response_json
                
        except httpx.RequestError as network_e:
            # Capturar errores de red específicos con Anthropic y lanzarlo para que routes.py mande HTTP 500
            raise Exception(f"Error de red al contactar con la IA de Anthropic: {str(network_e)}")
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {"estado": "INCOMPLETO", "mensaje_usuario": f"Lo siento, ocurrió un error interno al generar el reporte: {str(e)}"}

    async def _generar_reporte_fisico(self, response_json: Dict[str, Any]):
        parametros = response_json.get("parametros_query", {})
        filtros = parametros.get("filtros", {})
        formatos = parametros.get("formatos", [])
        graficos = parametros.get("graficos", [])
        if not formatos and parametros.get("formato"):
            formatos = [parametros.get("formato").upper()]
        if not formatos:
            formatos = ["EXCEL"]
            
        formatos = [f.upper() for f in formatos]
        
        agrupacion = parametros.get("agrupacion", "TRAMITE")
        
        # 1. Obtener datos de Spring Boot
        async with httpx.AsyncClient(timeout=60.0) as client:
            try:
                # Determinar el endpoint adecuado según la agrupación
                if agrupacion in ["DEPARTAMENTO", "FUNCIONARIO"]:
                    endpoint_url = self.backend_url.replace("/tramites/report-data", "/analytics/report-data")
                    filtros["agrupacion"] = agrupacion
                    resp = await client.post(endpoint_url, json=filtros)
                else:
                    resp = await client.post(self.backend_url, json=filtros)
                    
                resp.raise_for_status()
                datos_backend = resp.json()
            except httpx.RequestError as exc:
                raise Exception(f"Fallo de conexión de red con el Backend Spring Boot al solicitar {exc.request.url}.")
            except httpx.HTTPStatusError as exc:
                raise Exception(f"El Backend devolvió un error HTTP {exc.response.status_code}.")
            except Exception as e:
                raise Exception(f"No se pudo conectar con el Backend Spring Boot. Error interno: {str(e)}")
            
        if not datos_backend:
            datos_backend = [{"Mensaje": "No se encontraron resultados para los filtros especificados."}]
            
        # 2. Convertir a DataFrame dinámicamente
        df_list = []
        for row in datos_backend:
            if "Mensaje" in row:
                df_list.append(row)
            elif agrupacion == "TRAMITE":
                df_list.append({
                    "ID Tramite": row.get("id"),
                    "Nombre Plantilla": row.get("nombrePlantilla"),
                    "Estado": row.get("estadoGlobal"),
                    "Prioridad": row.get("prioridad"),
                    "Cliente Email": row.get("clienteEmail"),
                    "Fecha Creacion": row.get("fechaCreacion"),
                    "Fecha Finalizacion": row.get("fechaFinalizacion")
                })
            else:
                # Para DEPARTAMENTO o FUNCIONARIO tomamos los campos devueltos dinámicamente
                mapped_row = {}
                for key, val in row.items():
                    # Normalizar llaves para excel
                    capitalized_key = str(key).replace("_", " ").title()
                    mapped_row[capitalized_key] = val
                df_list.append(mapped_row)
        
        df = pd.DataFrame(df_list)
        
        # Generar gráficos si fueron solicitados
        graficos_bytes = []
        if graficos and not df.empty and agrupacion in ["DEPARTAMENTO", "FUNCIONARIO"]:
            graficos_bytes = self._generar_graficos(df, graficos)
            
        generated_files = [] # List of tuples (filename, content_bytes)

        for formato in formatos:
            buf = BytesIO()
            if formato == "EXCEL":
                # Escribimos a un buffer temporal para cargarlo con openpyxl
                temp_buf = BytesIO()
                df.to_excel(temp_buf, index=False, engine='openpyxl')
                temp_buf.seek(0)
                
                # Aplicamos estilos elegantes
                wb = openpyxl.load_workbook(temp_buf)
                ws = wb.active
                ws.title = "Reporte Analítico"
                
                header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
                header_font = Font(color="FFFFFF", bold=True, size=11)
                
                for cell in ws[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                for col in ws.columns:
                    max_length = 0
                    col_letter = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    ws.column_dimensions[col_letter].width = min(max_length + 2, 45)
                    
                wb.save(buf)
                
                if graficos_bytes:
                    for i, g_buf in enumerate(graficos_bytes):
                        img = OpenpyxlImage(g_buf)
                        start_row = len(df) + 3 + (i * 25)
                        ws.add_image(img, f"A{start_row}")
                    buf.seek(0)
                    wb.save(buf)

                generated_files.append(("reporte_dinamico.xlsx", buf.getvalue()))
                
            elif formato == "PDF":
                doc = SimpleDocTemplate(buf, pagesize=landscape(letter), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
                elements = []
                styles = getSampleStyleSheet()
                title_style = styles['Title']
                title_style.textColor = colors.HexColor("#1F4E78")
                
                elements.append(Paragraph("Reporte Dinámico de Trámites", title_style))
                elements.append(Paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
                elements.append(Spacer(1, 20))
                
                data = [df.columns.values.tolist()] + df.values.tolist()
                data = [[str(item)[:40] + '...' if len(str(item)) > 40 else str(item) for item in row] for row in data]
                
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1F4E78")),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 10),
                    ('BOTTOMPADDING', (0,0), (-1,0), 10),
                    ('TOPPADDING', (0,0), (-1,0), 10),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor("#FFFFFF"), colors.HexColor("#F4F6F9")]),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#DDDDDD")),
                    ('FONTSIZE', (0,1), (-1,-1), 8),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ]))
                elements.append(table)
                
                from reportlab.platypus import Image as RLImage
                if graficos_bytes:
                    for g_buf in graficos_bytes:
                        elements.append(Spacer(1, 20))
                        # g_buf contains the image data. RLImage can read from a file-like object, but we need to reset the pointer.
                        g_buf.seek(0)
                        elements.append(RLImage(g_buf, width=500, height=312))

                doc.build(elements)
                generated_files.append(("reporte_dinamico.pdf", buf.getvalue()))
                
            elif formato == "WORD":
                doc = docx.Document()
                title = doc.add_heading('Reporte Dinámico de Trámites', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                p = doc.add_paragraph(f"Generado automáticamente el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                t = doc.add_table(rows=1, cols=len(df.columns))
                t.style = 'Light Shading Accent 1' # Estilo integrado elegante en Word
                
                hdr_cells = t.rows[0].cells
                for i, column in enumerate(df.columns):
                    hdr_cells[i].text = str(column)
                    
                for index, row in df.iterrows():
                    row_cells = t.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)
                        
                if graficos_bytes:
                    for g_buf in graficos_bytes:
                        doc.add_paragraph("")
                        g_buf.seek(0)
                        doc.add_picture(g_buf, width=Inches(6.0))

                doc.save(buf)
                generated_files.append(("reporte_dinamico.docx", buf.getvalue()))
                
        # Si no se generó nada, fallback a Excel
        if not generated_files:
            df.to_excel(buf, index=False, engine='openpyxl')
            generated_files.append(("reporte_dinamico.xlsx", buf.getvalue()))

        # Retornar ZIP si hay más de 1 archivo, sino el archivo individual
        final_buffer = BytesIO()
        if len(generated_files) > 1:
            with zipfile.ZipFile(final_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for fname, fcontent in generated_files:
                    zipf.writestr(fname, fcontent)
            final_buffer.seek(0)
            return final_buffer, "reportes_multiples.zip", "application/zip"
        else:
            fname, fcontent = generated_files[0]
            final_buffer.write(fcontent)
            final_buffer.seek(0)
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            if fname.endswith(".pdf"):
                mime = "application/pdf"
            elif fname.endswith(".docx"):
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return final_buffer, fname, mime

    def _generar_graficos(self, df: pd.DataFrame, graficos: List[str]) -> List[BytesIO]:
        graficos_generados = []
        if df.empty or len(df.columns) < 2:
            return graficos_generados
            
        label_col = df.columns[1] if len(df.columns) > 1 and "Nombre" in str(df.columns[1]) else df.columns[0]
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) == 0:
            return graficos_generados
        
        val_col = numeric_cols[0] 
        
        for tipo in graficos:
            tipo = tipo.upper()
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if tipo == "PASTEL":
                ax.pie(df[val_col], labels=df[label_col].astype(str).str[:15], autopct='%1.1f%%', startangle=90)
                ax.set_title(f"{val_col} por {label_col}")
            elif tipo == "LINEAS":
                ax.plot(df[label_col].astype(str).str[:15], df[val_col], marker='o', color='#1F4E78')
                ax.set_ylabel(val_col)
                ax.set_title(f"{val_col} por {label_col}")
                plt.xticks(rotation=45, ha='right')
            else:
                ax.bar(df[label_col].astype(str).str[:15], df[val_col], color='#1F4E78')
                ax.set_ylabel(val_col)
                ax.set_title(f"{val_col} por {label_col}")
                plt.xticks(rotation=45, ha='right')
                
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png')
            plt.close(fig)
            buf.seek(0)
            graficos_generados.append(buf)
            
        return graficos_generados
